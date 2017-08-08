from io import BytesIO

from docx import Document

from django.http import HttpResponse
from django.urls import reverse_lazy
from django.views import View
from django.views.generic import CreateView
from django.views.generic import ListView
from django.views.generic.detail import SingleObjectMixin
from docx.shared import Pt

from agreements.forms import AgreementCreateForm
from agreements.utils import prepare_values
from cards.models import Request
from .models import Agreement


class AgreementCreateView(CreateView):
    model = Agreement
    form_class = AgreementCreateForm
    success_url = reverse_lazy('agreement_list')
    template_name = 'agreement/create.html'

    def get_form_kwargs(self):
        kwargs = super().get_form_kwargs()
        request = None
        if self.kwargs.get('pk'):
            try:
                request = Request.objects.get(pk=self.kwargs.get('pk'))
            except Request.DoesNotExist:
                pass

        if request:
            kwargs['initial'] = {
                'seller_user_last_name': request.user_last_name,
                'seller_user_first_name': request.user_first_name,
                'seller_user_middle_name': request.user_middle_name,
                'ts_vin': request.ts_vin,
                'ts_body_num': request.ts_body_num,
                'ts_frame_num': request.ts_frame_num,
                'ts_mark': request.ts_mark,
                'ts_model': request.ts_model,
                'ts_year': request.ts_year,
                'ts_reg_num': request.ts_reg_num,
            }
            if request.doc_type == 1:
                kwargs['initial']['srts_serial'] = request.doc_serial
                kwargs['initial']['srts_num'] = request.doc_num
                kwargs['initial']['srts_issued_date'] = request.doc_issued_date
                kwargs['initial']['srts_issued_by'] = request.doc_issued_by
            elif request.doc_type == 1:
                kwargs['initial']['pts_serial'] = request.doc_serial
                kwargs['initial']['pts_num'] = request.doc_num
                kwargs['initial']['pts_issued_date'] = request.doc_issued_date
                kwargs['initial']['pts_issued_by'] = request.doc_issued_by

        return kwargs

    def form_valid(self, form):
        """
        If the form is valid
        """
        self.object = form.save(commit=False)
        self.object.author_id = self.request.user.pk
        return super().form_valid(form)


class AgreementListView(ListView):
    template_name = 'agreement/list.html'
    paginate_by = 30

    def get_queryset(self):
        return Agreement.objects.filter(author_id=self.request.user.pk)


class AgreementDownloadView(SingleObjectMixin, View):
    def get_queryset(self):
        return Agreement.objects.filter(author_id=self.request.user.pk)

    def get(self, request, *args, **kwargs):
        self.object = self.get_object()
        doc_values = prepare_values(self.object)

        if self.object.is_repr:
            template_file = "./agreements/doc_templates/dkp_repr_base.docx"
        else:
            template_file = "./agreements/doc_templates/dkp_base.docx"

        document = Document(template_file)

        style = document.styles['Normal']
        font = style.font
        font.name = 'Times New Roman'
        font.size = Pt(8)

        for paragraph in document.paragraphs:
            for field, value in doc_values.items():
                expr = '{{' + field + '}}'
                if expr in paragraph.text:
                    paragraph.text = paragraph.text.replace(expr, str(value))
                    if field != 'pk':
                        paragraph.style = style

        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for field, value in doc_values.items():
                            expr = '{{' + field + '}}'
                            if expr in paragraph.text:
                                paragraph.text = paragraph.text.replace(expr, str(value))
                                paragraph.style = style

        f = BytesIO()
        document.save(f)
        length = f.tell()
        f.seek(0)

        response = HttpResponse(
            f.getvalue(),
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        response['Content-Disposition'] = 'filename=dkp_{}.docx'.format(self.object.pk)
        response['Content-Length'] = length
        return response